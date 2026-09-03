"""Deliverable engine tests.

The headline output of this platform is a filed document, so every writer is
exercised for real: a file is produced, opened again, and checked for the
content and provenance a reviewer needs.

These exist because a wrong keyword in one python-docx call meant no approval
note was ever produced, and nothing caught it until a full run was inspected
by hand.
"""

from __future__ import annotations

from datetime import datetime, timezone

import pytest

from backend.core.schemas import (
    EvidenceItem,
    ModelRole,
    RoutingDecision,
    Sensitivity,
    VerificationCheck,
    VerificationReport,
)
from backend.tools.deliverables import DeliverableEngine

TASK_ID = "test-task-0001"

CONTENT = {
    "title": "Approval Note — Vessel V-2104",
    "reference": "INS-2026-0417",
    "summary": "Wall loss at shell course 2 is the governing location [S1].",
    "sections": [
        {
            "heading": "Findings",
            "body": "Measured thickness 9.4 mm against a minimum of 6.0 mm [S1].",
            "bullets": ["Cladding damaged over 35% of course 2"],
        }
    ],
    "findings": [
        {
            "description": "Corrosion under insulation at shell course 2",
            "severity": "medium",
            "reference": "SOP-INS-014 Clause 5",
        }
    ],
    "recommendation": "Repair cladding during the current shutdown window.",
    "approval_statement": "Approval is sought to continue service until the next campaign.",
}

EVIDENCE = [
    EvidenceItem(
        id="S1",
        source_document="SOP-INS-014 — Pressure Vessel Inspection",
        location="section: 4. Corrosion Rate and Remaining Life",
        excerpt="Remaining life shall be calculated as (t_current - t_min) / rate.",
        classification=Sensitivity.CONFIDENTIAL,
        version="4.2",
        department="inspection",
    )
]

ROUTING = [
    RoutingDecision(
        requested_role=ModelRole.VISION,
        required_capabilities=["vision"],
        selected_model="qwen2.5vl:3b",
        selected_display_name="Qwen2.5-VL 3B",
        rule="visual_understanding",
        reason="scan requires a vision model",
        decided_at=datetime.now(timezone.utc),
    )
]

VERIFICATION = VerificationReport(
    valid=True,
    checks=[
        VerificationCheck(
            name="source_verification",
            kind="source",
            passed=True,
            detail="2 of 2 material claims supported.",
        )
    ],
    material_claims_total=2,
    material_claims_supported=2,
    completed_at=datetime.now(timezone.utc),
)

CALCULATIONS = [
    {
        "label": "corrosion rate",
        "expression": "(12.0 - 9.4) / 4.0",
        "expected": 0.65,
        "recomputed": 0.65,
        "units": "mm/year",
        "matched": True,
    }
]


@pytest.fixture
def engine() -> DeliverableEngine:
    return DeliverableEngine()


class TestWordDeliverable:
    def test_produces_a_readable_file(self, engine: DeliverableEngine) -> None:
        deliverable = engine.render_docx(
            task_id=TASK_ID,
            content=CONTENT,
            evidence=EVIDENCE,
            routing=ROUTING,
            verification=VERIFICATION,
            author="Integrity Engineer",
            calculations=CALCULATIONS,
        )
        assert deliverable.format == "docx"
        assert deliverable.size_bytes > 0
        assert len(deliverable.sha256) == 64
        assert not deliverable.released, "a fresh deliverable is held until approved"

    def test_document_carries_its_evidence_and_provenance(
        self, engine: DeliverableEngine
    ) -> None:
        """A reviewer must be able to trace the note without leaving it."""
        import docx

        from backend.core.config import get_config

        deliverable = engine.render_docx(
            task_id=TASK_ID,
            content=CONTENT,
            evidence=EVIDENCE,
            routing=ROUTING,
            verification=VERIFICATION,
            author="Integrity Engineer",
            calculations=CALCULATIONS,
        )
        path = get_config().settings.path("deliverables") / TASK_ID / deliverable.filename
        document = docx.Document(str(path))

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        tables = "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        everything = f"{text}\n{tables}"

        assert "Approval Note" in everything.title() or CONTENT["title"] in everything
        assert "S1" in everything, "citation markers must survive into the document"
        assert "SOP-INS-014" in everything, "the cited source must be named"
        assert CONTENT["recommendation"] in everything
        assert "Qwen2.5-VL 3B" in everything, "the model used must be recorded"
        assert "0.65" in everything, "the recomputed figure must appear"
        assert "no external calls" in everything.lower()

    def test_handles_minimal_content_without_crashing(
        self, engine: DeliverableEngine
    ) -> None:
        deliverable = engine.render_docx(
            task_id=TASK_ID,
            content={"title": "Bare note", "sections": []},
            evidence=[],
            routing=[],
            verification=None,
            author="Operator",
        )
        assert deliverable.size_bytes > 0


class TestOtherFormats:
    def test_workbook(self, engine: DeliverableEngine) -> None:
        from openpyxl import load_workbook

        from backend.core.config import get_config

        deliverable = engine.render_xlsx(
            task_id=TASK_ID,
            content=CONTENT,
            evidence=EVIDENCE,
            tables=[{"name": "Readings", "headers": ["Location", "2026"], "rows": [["c2", 9.4]]}],
        )
        path = get_config().settings.path("deliverables") / TASK_ID / deliverable.filename
        workbook = load_workbook(str(path))
        assert "Summary" in workbook.sheetnames
        assert "Readings" in workbook.sheetnames
        assert "Evidence" in workbook.sheetnames

    def test_presentation(self, engine: DeliverableEngine) -> None:
        from pptx import Presentation

        from backend.core.config import get_config

        deliverable = engine.render_pptx(
            task_id=TASK_ID, content=CONTENT, evidence=EVIDENCE
        )
        path = get_config().settings.path("deliverables") / TASK_ID / deliverable.filename
        assert len(Presentation(str(path)).slides) >= 4

    def test_markdown_keeps_citations(self, engine: DeliverableEngine) -> None:
        from backend.core.config import get_config

        deliverable = engine.render_markdown(
            task_id=TASK_ID, content=CONTENT, evidence=EVIDENCE
        )
        path = get_config().settings.path("deliverables") / TASK_ID / deliverable.filename
        body = path.read_text(encoding="utf-8")
        assert "[S1]" in body
        assert "SOP-INS-014" in body

    def test_every_declared_format_renders(self, engine: DeliverableEngine) -> None:
        """Whatever the analyzer may select, the engine must be able to write."""
        for fmt in ("docx", "xlsx", "pptx", "md"):
            deliverable = engine.render(
                fmt,
                task_id=TASK_ID,
                content=CONTENT,
                evidence=EVIDENCE,
                routing=ROUTING,
                verification=VERIFICATION,
            )
            assert deliverable.size_bytes > 0, fmt

    def test_unsupported_format_is_refused_clearly(self, engine: DeliverableEngine) -> None:
        with pytest.raises(ValueError, match="Unsupported deliverable format"):
            engine.render("pdf-with-signatures", task_id=TASK_ID, content=CONTENT)
